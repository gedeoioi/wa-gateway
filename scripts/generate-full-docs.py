from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
for level in range(1, 4):
    doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0x07, 0x5E, 0x54)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F0F0")
    run._element.get_or_add_rPr().append(shading)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

title = doc.add_heading("WA Gateway", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WhatsApp Gateway Enterprise")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x07, 0x5E, 0x54)
run.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Panduan Lengkap Deploy & Setting")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Versi 1.0.0 | Agustus 2026")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_page_break()

doc.add_heading("Daftar Isi", level=1)
toc = [
    "Bagian 1: Persiapan (VPS, Domain, Tools)",
    "Bagian 2: Install Dependencies di VPS",
    "Bagian 3: Setup Database",
    "Bagian 4: Deploy Aplikasi",
    "Bagian 5: Konfigurasi Nginx",
    "Bagian 6: Menghubungkan Domain & SSL",
    "Bagian 7: Setup Firewall",
    "Bagian 8: Penggunaan Aplikasi",
    "Bagian 9: API Documentation",
    "Bagian 10: Maintenance & Backup",
    "Bagian 11: Troubleshooting",
    "Bagian 12: Arsitektur & Tech Stack",
]
for item in toc:
    doc.add_paragraph(item, style="List Number")
doc.add_page_break()

# === BAGIAN 1 ===
doc.add_heading("Bagian 1: Persiapan", level=1)
doc.add_heading("1.1 Spesifikasi VPS", level=2)
add_table(doc, ["Resource", "Minimum", "Recommended"],
    [["OS", "Ubuntu 22.04 LTS", "Ubuntu 24.04 LTS"],
     ["CPU", "1 vCPU", "2 vCPU"],
     ["RAM", "1 GB", "2 GB"],
     ["Storage", "20 GB SSD", "40 GB SSD"]])

doc.add_heading("1.2 Persiapan Domain", level=2)
doc.add_paragraph("Beli domain dari registrar Indonesia (Niagahoster, Domainesia, Rumahweb) atau internasional (Cloudflare, Namecheap).")

doc.add_heading("1.3 Tools", level=2)
add_table(doc, ["Tool", "Fungsi"],
    [["SSH Client", "Akses VPS (PuTTY / Terminal)"],
     ["Git", "Version control"],
     ["DBeaver / HeidiSQL", "Database GUI (opsional)"]])
doc.add_page_break()

# === BAGIAN 2 ===
doc.add_heading("Bagian 2: Install Dependencies di VPS", level=1)
doc.add_paragraph("SSH ke VPS: ssh root@IP_VPS")

doc.add_heading("2.1 Update System", level=2)
add_code(doc, "sudo apt update && sudo apt upgrade -y")

doc.add_heading("2.2 Install Node.js 22", level=2)
add_code(doc, "curl -fsSL https://deb.nodesource.com/setup_22.x | sudo -E bash -\nsudo apt install -y nodejs\nnode -v && npm -v")

doc.add_heading("2.3 Install MySQL", level=2)
add_code(doc, "sudo apt install -y mysql-server\nsudo systemctl start mysql && sudo systemctl enable mysql\nsudo mysql_secure_installation")

doc.add_heading("2.4 Install Redis", level=2)
add_code(doc, "sudo apt install -y redis-server\nsudo systemctl start redis-server && sudo systemctl enable redis-server\nredis-cli ping  # harus: PONG")

doc.add_heading("2.5 Install Nginx & PM2", level=2)
add_code(doc, "sudo apt install -y nginx\nsudo systemctl start nginx && sudo systemctl enable nginx\nsudo npm install -g pm2")
doc.add_page_break()

# === BAGIAN 3 ===
doc.add_heading("Bagian 3: Setup Database", level=1)
doc.add_heading("3.1 Buat Database & User", level=2)
add_code(doc, "sudo mysql")
add_code(doc, "CREATE DATABASE wa_gateway CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;\nCREATE USER 'wa_user'@'localhost' IDENTIFIED BY 'PASSWORD_KUAT';\nGRANT ALL PRIVILEGES ON wa_gateway.* TO 'wa_user'@'localhost';\nFLUSH PRIVILEGES;\nEXIT;")
add_note(doc, "Ganti PASSWORD_KUAT dengan password yang kuat dan unik!")

doc.add_heading("3.2 Generate Secret Keys", level=2)
add_code(doc, "node -e \"console.log(require('crypto').randomBytes(32).toString('hex'))\"")
doc.add_paragraph("Jalankan 3 kali untuk JWT_SECRET, JWT_REFRESH_SECRET, ENCRYPTION_KEY.")
doc.add_page_break()

# === BAGIAN 4 ===
doc.add_heading("Bagian 4: Deploy Aplikasi", level=1)
doc.add_heading("4.1 Clone Repository", level=2)
add_code(doc, "cd /var/www\ngit clone https://github.com/yourusername/wa-gateway.git\nsudo chown -R $USER:$USER wa-gateway\ncd wa-gateway")

doc.add_heading("4.2 Install & Setup", level=2)
add_code(doc, "npm install\ncp .env.example .env\nnano .env")
doc.add_paragraph("Isi .env:")
add_code(doc, "NODE_ENV=production\nDATABASE_URL=\"mysql://wa_user:PASSWORD@localhost:3306/wa_gateway\"\nREDIS_URL=\"redis://localhost:6379\"\nJWT_SECRET=hasil_generate\nJWT_REFRESH_SECRET=hasil_generate\nJWT_EXPIRES_IN=15m\nJWT_REFRESH_EXPIRES_IN=7d\nENCRYPTION_KEY=hasil_generate\nBACKEND_PORT=3001\nFRONTEND_URL=https://domain-anda.com\nBACKEND_URL=https://domain-anda.com\nNEXT_PUBLIC_API_URL=\nNEXT_PUBLIC_SOCKET_URL=")

doc.add_heading("4.3 Build & Database", level=2)
add_code(doc, "npm run db:generate\nnpm run db:migrate\nnpm run db:seed\nnpm run build")

doc.add_heading("4.4 Jalankan dengan PM2", level=2)
add_code(doc, "pm2 start ecosystem.config.js --env production\npm2 save\npm2 startup\npm2 status")
add_table(doc, ["Name", "Status", "Port"],
    [["wa-gateway-backend", "online", "3001"],
     ["wa-gateway-frontend", "online", "3000"]])
doc.add_page_break()

# === BAGIAN 5 ===
doc.add_heading("Bagian 5: Konfigurasi Nginx", level=1)
add_code(doc, "sudo nano /etc/nginx/sites-available/wa-gateway")
doc.add_paragraph("Isi dengan config berikut:")
add_code(doc, "server {\n    listen 80;\n    server_name domain-anda.com www.domain-anda.com;\n\n    location / {\n        proxy_pass http://127.0.0.1:3000;\n        proxy_http_version 1.1;\n        proxy_set_header Upgrade \$http_upgrade;\n        proxy_set_header Connection \"upgrade\";\n        proxy_set_header Host \$host;\n        proxy_set_header X-Real-IP \$remote_addr;\n        proxy_set_header X-Forwarded-For \$proxy_add_x_forwarded_for;\n        proxy_set_header X-Forwarded-Proto \$scheme;\n    }\n\n    location /api/ {\n        proxy_pass http://127.0.0.1:3001;\n        proxy_http_version 1.1;\n        proxy_set_header Upgrade \$http_upgrade;\n        proxy_set_header Connection \"upgrade\";\n        proxy_set_header Host \$host;\n        proxy_set_header X-Real-IP \$remote_addr;\n        proxy_read_timeout 300s;\n    }\n\n    location /socket.io/ {\n        proxy_pass http://127.0.0.1:3001;\n        proxy_http_version 1.1;\n        proxy_set_header Upgrade \$http_upgrade;\n        proxy_set_header Connection \"upgrade\";\n    }\n\n    location /uploads/ {\n        proxy_pass http://127.0.0.1:3001/uploads/;\n    }\n\n    location /api-docs {\n        proxy_pass http://127.0.0.1:3001;\n    }\n\n    client_max_body_size 16M;\n}")

add_code(doc, "sudo ln -s /etc/nginx/sites-available/wa-gateway /etc/nginx/sites-enabled/\nsudo rm /etc/nginx/sites-enabled/default\nsudo nginx -t\nsudo systemctl restart nginx")
doc.add_page_break()

# === BAGIAN 6 ===
doc.add_heading("Bagian 6: Domain & SSL", level=1)
doc.add_heading("6.1 DNS Record", level=2)
add_table(doc, ["Type", "Name", "Value"],
    [["A", "@", "IP_VPS"],
     ["A", "www", "IP_VPS"]])
doc.add_paragraph("Tunggu 5-30 menit propagasi. Verifikasi: nslookup domain-anda.com")

doc.add_heading("6.2 Update .env & Rebuild", level=2)
add_code(doc, "nano .env  # set FRONTEND_URL dan BACKEND_URL ke domain\nnpm run build:frontend\npm2 restart all")

doc.add_heading("6.3 Install SSL", level=2)
add_code(doc, "sudo apt install -y certbot python3-certbot-nginx\nsudo certbot --nginx -d domain-anda.com -d www.domain-anda.com\nsudo certbot renew --dry-run")

doc.add_heading("6.4 Cloudflare (Opsional)", level=2)
doc.add_paragraph("1. Buat akun cloudflare.com\n2. Tambahkan domain\n3. Ganti nameserver di registrar ke Cloudflare\n4. Tambah A record di Cloudflare DNS\n5. Set proxy Proxied (orange cloud)\n6. SSL/TLS mode: Full (Strict)")
doc.add_page_break()

# === BAGIAN 7 ===
doc.add_heading("Bagian 7: Firewall", level=1)
add_code(doc, "sudo ufw allow OpenSSH\nsudo ufw allow 'Nginx Full'\nsudo ufw enable")
doc.add_page_break()

# === BAGIAN 8 ===
doc.add_heading("Bagian 8: Penggunaan Aplikasi", level=1)

doc.add_heading("8.1 Login Pertama", level=2)
add_table(doc, ["Field", "Value"], [["Username", "admin"], ["Password", "admin123"]])
add_note(doc, "Segera ganti password! Settings > Security > Change Password.")

doc.add_heading("8.2 Kustomisasi Tampilan (Admin)", level=2)
doc.add_paragraph("Menu Appearance:")
add_table(doc, ["Setting", "Fungsi"],
    [["Site Name", "Nama website"],
     ["Logo", "Upload logo (sidebar & login)"],
     ["Favicon", "Upload favicon (tab browser)"],
     ["Colors", "8 preset warna + custom color picker"],
     ["Footer Text", "Teks footer"]])

doc.add_heading("8.3 Manajemen User (Admin)", level=2)
add_table(doc, ["Role", "Akses"],
    [["Admin", "Semua fitur + kelola user + kustomisasi"],
     ["Operator", "Kirim pesan, broadcast, device, kontak"],
     ["Viewer", "Read-only"]])

doc.add_heading("8.4 Menghubungkan WhatsApp", level=2)
add_code(doc, "1. Menu Devices > Add Device > beri nama\n2. Klik Connect > tunggu QR code\n3. Buka WhatsApp HP > Settings > Linked Devices\n4. Scan QR > tunggu Connected")

doc.add_heading("8.5 Mengirim Pesan", level=2)
add_code(doc, "1. Menu Messages > Send Message\n2. Pilih Device > masukkan nomor (628xxx)\n3. Ketik pesan > Send")

doc.add_heading("8.6 Broadcast", level=2)
add_code(doc, "1. Menu Broadcasts > New Broadcast\n2. Pilih Device > nama campaign > pesan\n3. Masukkan nomor (satu per baris)\n4. Create > Start > monitor progress")

doc.add_heading("8.7 Keamanan", level=2)
add_table(doc, ["Fitur", "Detail"],
    [["Change Password", "Ganti password dengan verifikasi"],
     ["Login History", "10 aktivitas terakhir (IP, device)"],
     ["Revoke Sessions", "Logout semua device lain"],
     ["Account Lockout", "5x gagal > terkunci 15 menit"]])
doc.add_page_break()

# === BAGIAN 9 ===
doc.add_heading("Bagian 9: API Documentation", level=1)
doc.add_paragraph("Swagger UI: https://domain-anda.com/api-docs")
add_table(doc, ["Tag", "Endpoints"],
    [["Auth", "8 endpoints"],
     ["Users", "4 endpoints (admin)"],
     ["Devices", "8 endpoints"],
     ["Messages", "5 endpoints"],
     ["Broadcasts", "6 endpoints"],
     ["Contacts", "5 endpoints"],
     ["Contact Groups", "4 endpoints"],
     ["Upload", "1 endpoint"],
     ["Site Settings", "2 endpoints"]])
doc.add_paragraph("Total: 44 endpoints terdokumentasi")
doc.add_page_break()

# === BAGIAN 10 ===
doc.add_heading("Bagian 10: Maintenance", level=1)
doc.add_heading("10.1 Update Aplikasi", level=2)
add_code(doc, "cd /var/www/wa-gateway\ngit pull\nnpm install\nnpm run db:generate\nnpx prisma db push\nnpm run build\npm2 restart all")

doc.add_heading("10.2 Backup Database", level=2)
add_code(doc, "# Backup\nmysqldump -u wa_user -p wa_gateway > backup_\$(date +%Y%m%d).sql\n\n# Restore\nmysql -u wa_user -p wa_gateway < backup_20260808.sql\n\n# Cron job otomatis (jam 2 pagi)\ncrontab -e\n0 2 * * * mysqldump -u wa_user -pPASS wa_gateway > /var/backups/wa_\$(date +\\%Y\\%m\\%d).sql")

doc.add_heading("10.3 Monitoring", level=2)
add_code(doc, "pm2 status\npm2 monit\npm2 logs --err\nsudo ss -tlnp | grep -E '3000|3001|3306|6379'\ndf -h\nfree -m")
doc.add_page_break()

# === BAGIAN 11 ===
doc.add_heading("Bagian 11: Troubleshooting", level=1)
add_table(doc, ["Masalah", "Solusi"],
    [["502 Bad Gateway", "pm2 restart wa-gateway-backend"],
     ["Failed to fetch", "npm run build:frontend, cek FRONTEND_URL"],
     ["Too many requests", "Restart backend"],
     ["Session expired", "Login ulang"],
     ["QR code tidak muncul", "Klik Connect, tunggu beberapa detik"],
     ["Upload gagal", "mkdir -p apps/backend/uploads"],
     ["Logo tidak muncul", "Cek Nginx /uploads/ proxy"],
     ["Prisma error", "GRANT ALL ON *.* TO user@localhost"],
     ["PM2 crash", "npm run build:backend, pm2 logs"],
     ["SSL gagal", "Cek DNS, tunggu propagate, port 80 terbuka"],
     ["Port digunakan", "sudo kill \$(sudo lsof -t -i:3001)"],
     ["Redis down", "sudo systemctl restart redis-server"],
     ["MySQL down", "sudo systemctl restart mysql"]])
doc.add_page_break()

# === BAGIAN 12 ===
doc.add_heading("Bagian 12: Arsitektur", level=1)
add_code(doc, "Internet\n    |\n    v\n+-----------+\n|  Nginx    | :80 -> :443\n+-----+-----+\n      |\n      +-- /          -> Next.js (3000)\n      +-- /api/      -> Express (3001)\n      +-- /socket.io -> Express (3001)\n      +-- /uploads/  -> Express (3001)\n                            |\n                       +----+----+\n                       |         |\n                    MySQL     Redis\n                   (3306)    (6379)")

doc.add_heading("Tech Stack", level=2)
add_table(doc, ["Layer", "Technology"],
    [["Frontend", "Next.js 15, React 19, Tailwind CSS, TypeScript"],
     ["Backend", "Express.js, TypeScript, Prisma ORM"],
     ["Database", "MySQL 8.0"],
     ["Cache/Queue", "Redis + BullMQ"],
     ["WhatsApp", "Baileys"],
     ["Auth", "JWT (access + refresh token)"],
     ["Process Manager", "PM2"],
     ["Reverse Proxy", "Nginx"]])

doc.add_heading("Fitur Utama", level=2)
add_table(doc, ["Fitur", "Detail"],
    [["Multi-device WhatsApp", "Beberapa nomor WhatsApp"],
     ["Pesan & Broadcast", "Teks, gambar, video, dokumen, audio"],
     ["Multi-user", "Admin, Operator, Viewer"],
     ["Keamanan", "Lockout, login history, session management"],
     ["Kustomisasi", "Logo, favicon, warna tema"],
     ["API Docs", "Swagger UI, 44 endpoints"],
     ["Real-time", "WebSocket untuk status & progress"],
     ["Responsive", "Desktop, tablet, mobile"]])

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WA Gateway Enterprise v1.0.0 - Panduan Lengkap Deploy & Setting")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save("D:/wa-gateway/docs/Deploy_Setting_WA_Gateway.docx")
print("Document saved!")
