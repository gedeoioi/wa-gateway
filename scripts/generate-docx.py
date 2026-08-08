from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    heading_style = doc.styles[f"Heading {level}"]
    heading_style.font.color.rgb = RGBColor(0x07, 0x5E, 0x54)

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F0F0")
    run._element.get_or_add_rPr().append(shading)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# TITLE
title = doc.add_heading("WA Gateway - WhatsApp Gateway Enterprise", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Panduan Deploy ke VPS")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# DAFTAR ISI
doc.add_heading("Daftar Isi", level=1)
toc_items = [
    "1. Spesifikasi Minimum VPS",
    "2. Install Dependencies",
    "3. Setup MySQL",
    "4. Clone & Setup Project",
    "5. Build & Setup Database",
    "6. Setup PM2 (Process Manager)",
    "7. Setup Nginx (Reverse Proxy)",
    "8. Setup SSL (Lets Encrypt)",
    "9. Setup Firewall",
    "10. Commands Berguna",
    "11. Troubleshooting",
    "12. Arsitektur Deploy",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")

doc.add_page_break()

# 1
doc.add_heading("1. Spesifikasi Minimum VPS", level=1)
add_table(doc,
    ["Resource", "Minimum", "Recommended"],
    [
        ["OS", "Ubuntu 22.04 / 24.04 LTS", "Ubuntu 24.04 LTS"],
        ["CPU", "1 vCPU", "2 vCPU"],
        ["RAM", "1 GB", "2 GB"],
        ["Storage", "20 GB", "40 GB"],
    ]
)

# 2
doc.add_heading("2. Install Dependencies", level=1)
doc.add_paragraph("Jalankan perintah berikut di VPS untuk menginstall semua dependencies yang dibutuhkan:")
add_code_block(doc, """# Update system
sudo apt update && sudo apt upgrade -y

# Install Node.js 22
curl -fsSL https://deb.nodesource.com/setup_22.x | sudo -E bash -
sudo apt install -y nodejs

# Install MySQL
sudo apt install -y mysql-server
sudo systemctl start mysql
sudo systemctl enable mysql

# Install Redis
sudo apt install -y redis-server
sudo systemctl start redis-server
sudo systemctl enable redis-server

# Install Git
sudo apt install -y git

# Install PM2 (process manager)
sudo npm install -g pm2

# Install Nginx
sudo apt install -y nginx
sudo systemctl start nginx
sudo systemctl enable nginx

# Verifikasi
node -v && npm -v && mysql --version && redis-cli ping""")

# 3
doc.add_heading("3. Setup MySQL", level=1)
doc.add_paragraph("Buat database dan user untuk aplikasi:")
add_code_block(doc, "sudo mysql")
doc.add_paragraph("Jalankan SQL berikut di dalam MySQL shell:")
add_code_block(doc, """CREATE DATABASE wa_gateway CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;
CREATE USER 'wa_user'@'localhost' IDENTIFIED BY 'PASSWORD_KUAT_DISINI';
GRANT ALL PRIVILEGES ON wa_gateway.* TO 'wa_user'@'localhost';
FLUSH PRIVILEGES;
EXIT;""")
p = doc.add_paragraph()
run = p.add_run("Penting: ")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run("Ganti PASSWORD_KUAT_DISINI dengan password yang kuat dan unik.")

# 4
doc.add_heading("4. Clone & Setup Project", level=1)
doc.add_paragraph("Clone repository dan konfigurasi environment:")
add_code_block(doc, """# Clone repository
cd /var/www
sudo git clone https://github.com/yourusername/wa-gateway.git
sudo chown -R $USER:$USER wa-gateway
cd wa-gateway

# Install dependencies
npm install

# Setup environment
cp .env.example .env""")
doc.add_paragraph("Edit file .env:")
add_code_block(doc, "nano .env")
doc.add_paragraph("Isi file .env:")
add_code_block(doc, """NODE_ENV=production
LOG_LEVEL=info

DATABASE_URL="mysql://wa_user:PASSWORD_KUAT_DISINI@localhost:3306/wa_gateway"
REDIS_URL="redis://localhost:6379"

JWT_SECRET=<generate-random-secret>
JWT_REFRESH_SECRET=<generate-random-refresh-secret>
JWT_EXPIRES_IN=15m
JWT_REFRESH_EXPIRES_IN=7d

ENCRYPTION_KEY=<generate-64-hex-chars>

BACKEND_PORT=3001
FRONTEND_URL=https://domain-anda.com
BACKEND_URL=https://domain-anda.com
NEXT_PUBLIC_API_URL=https://domain-anda.com
NEXT_PUBLIC_SOCKET_URL=https://domain-anda.com""")
doc.add_paragraph("Generate secret key:")
add_code_block(doc, 'node -e "console.log(require(\'crypto\').randomBytes(32).toString(\'hex\'))"')
doc.add_paragraph("Jalankan 2 kali untuk JWT_SECRET dan JWT_REFRESH_SECRET yang berbeda.")

# 5
doc.add_heading("5. Build & Setup Database", level=1)
add_code_block(doc, """# Generate Prisma Client
npm run db:generate

# Jalankan migrasi
npm run db:migrate

# Seed database (buat admin user)
npm run db:seed

# Build untuk production
npm run build""")
p = doc.add_paragraph()
run = p.add_run("Default credentials: ")
run.bold = True
p.add_run("Username: admin | Password: admin123")
doc.add_paragraph("Segera ganti password setelah login pertama (Settings > Security).")

doc.add_page_break()

# 6
doc.add_heading("6. Setup PM2 (Process Manager)", level=1)
doc.add_paragraph("PM2 menjalankan aplikasi sebagai background process dengan auto-restart. File ecosystem.config.js sudah disediakan.")
add_code_block(doc, """# Jalankan dengan PM2
pm2 start ecosystem.config.js --env production

# Simpan process list
pm2 save

# Setup auto-start saat boot
pm2 startup
# Jalankan perintah yang diberikan output-nya

# Cek status
pm2 status""")
doc.add_paragraph("Verifikasi kedua proses berjalan:")
add_table(doc,
    ["Nama", "Status", "Port"],
    [
        ["wa-gateway-backend", "online", "3001"],
        ["wa-gateway-frontend", "online", "3000"],
    ]
)

# 7
doc.add_heading("7. Setup Nginx (Reverse Proxy)", level=1)
doc.add_paragraph("Buat konfigurasi Nginx:")
add_code_block(doc, "sudo nano /etc/nginx/sites-available/wa-gateway")
doc.add_paragraph("Isi dengan konfigurasi berikut:")
add_code_block(doc, """server {
    listen 80;
    server_name domain-anda.com;
    return 301 https://$server_name$request_uri;
}

server {
    listen 443 ssl http2;
    server_name domain-anda.com;

    ssl_certificate /etc/letsencrypt/live/domain-anda.com/fullchain.pem;
    ssl_certificate_key /etc/letsencrypt/live/domain-anda.com/privkey.pem;

    add_header X-Frame-Options "SAMEORIGIN" always;
    add_header X-Content-Type-Options "nosniff" always;
    add_header X-XSS-Protection "1; mode=block" always;

    # Frontend
    location / {
        proxy_pass http://127.0.0.1:3000;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection "upgrade";
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }

    # Backend API
    location /api/ {
        proxy_pass http://127.0.0.1:3001;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection "upgrade";
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
        proxy_read_timeout 300s;
    }

    # Socket.IO
    location /socket.io/ {
        proxy_pass http://127.0.0.1:3001;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection "upgrade";
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
    }

    # Swagger API Docs
    location /api-docs {
        proxy_pass http://127.0.0.1:3001;
    }

    client_max_body_size 10M;
}""")
doc.add_paragraph("Aktifkan site dan restart Nginx:")
add_code_block(doc, """sudo ln -s /etc/nginx/sites-available/wa-gateway /etc/nginx/sites-enabled/
sudo rm /etc/nginx/sites-enabled/default
sudo nginx -t
sudo systemctl restart nginx""")

# 8
doc.add_heading("8. Setup SSL (Lets Encrypt)", level=1)
add_code_block(doc, """# Install Certbot
sudo apt install -y certbot python3-certbot-nginx

# Generate SSL certificate
sudo certbot --nginx -d domain-anda.com

# Test auto-renewal
sudo certbot renew --dry-run""")

# 9
doc.add_heading("9. Setup Firewall", level=1)
add_code_block(doc, """sudo ufw allow OpenSSH
sudo ufw allow 'Nginx Full'
sudo ufw enable
sudo ufw status""")
add_table(doc,
    ["Port", "Protocol", "Service", "Status"],
    [
        ["22", "TCP", "SSH", "ALLOW"],
        ["80", "TCP", "HTTP (redirect to HTTPS)", "ALLOW"],
        ["443", "TCP", "HTTPS", "ALLOW"],
    ]
)

doc.add_page_break()

# 10
doc.add_heading("10. Commands Berguna", level=1)

doc.add_heading("PM2", level=2)
add_code_block(doc, """pm2 restart all
pm2 restart wa-gateway-backend
pm2 restart wa-gateway-frontend
pm2 logs
pm2 logs wa-gateway-backend --lines 50
pm2 monit""")

doc.add_heading("Update Project", level=2)
add_code_block(doc, """cd /var/www/wa-gateway
git pull
npm install
npm run build
pm2 restart all""")

doc.add_heading("Service Management", level=2)
add_code_block(doc, """sudo systemctl restart nginx
sudo systemctl restart mysql
sudo systemctl restart redis-server""")

# 11
doc.add_heading("11. Troubleshooting", level=1)

doc.add_heading("Cek Port", level=2)
add_code_block(doc, "sudo ss -tlnp | grep -E '3000|3001|3306|6379|80|443'")

doc.add_heading("Cek Status Service", level=2)
add_code_block(doc, """pm2 status
sudo systemctl status nginx
sudo systemctl status mysql
sudo systemctl status redis-server""")

doc.add_heading("Cek Logs Error", level=2)
add_code_block(doc, """pm2 logs --err
sudo tail -f /var/log/nginx/error.log""")

doc.add_heading("Masalah Umum", level=2)
add_table(doc,
    ["Masalah", "Solusi"],
    [
        ["Port sudah digunakan", "sudo kill $(sudo lsof -t -i:3001) lalu pm2 restart"],
        ["MySQL connection refused", "sudo systemctl restart mysql, cek DATABASE_URL"],
        ["Redis connection refused", "sudo systemctl restart redis-server, cek REDIS_URL"],
        ["502 Bad Gateway", "pm2 status, pm2 logs"],
        ["Prisma generate error", "pm2 stop all, npm run db:generate, pm2 restart all"],
        ["SSL expired", "sudo certbot renew && sudo systemctl restart nginx"],
    ]
)

# 12
doc.add_heading("12. Arsitektur Deploy", level=1)
add_code_block(doc, """Internet
    |
    v
+-----------+
|  Nginx    | :80 -> redirect :443
|  (SSL)    | :443
+-----+-----+
      |
      +-- /          -> Next.js Frontend (port 3000)
      +-- /api/      -> Express Backend  (port 3001)
      +-- /socket.io -> Express Backend  (port 3001)
                            |
                       +----+----+
                       |         |
                    MySQL     Redis
                   (3306)    (6379)""")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WA Gateway Enterprise - Panduan Deploy VPS")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save("D:/wa-gateway/docs/Deploy_VPS_WA_Gateway.docx")
print("Document saved!")
